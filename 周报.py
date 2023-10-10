from openpyxl import load_workbook     #Excel操作：用于加载已经存在的excel
from openpyxl import Workbook          #Excel操作：用于创建新的excel
from docx import Document              #Word操作：导入Docx基础包
from docx.shared import Cm,Inches,Pt   #Word操作：导入单位换算函数
from docx.oxml.ns import qn            #Word操作：中文字体模块
from docx.enum.text import WD_ALIGN_PARAGRAPH     #导入对齐选项
import time                            #导入时间，用来控制程序运行时间
import pandas as pd  
from datetime import datetime, timedelta 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX  
import re 
#路径
LoadPath="myLife.xlsx"      #加载excel路径(这里为相对路径，excel表与该程序在同一文件夹下就能识别，所以只用excel文件名即可)

#excel表格初始化
book=load_workbook(LoadPath)   #加载已有Excel文档
try:
    sheet=book['Activity']     #加载需要的工作簿（这里为excel表中的sheet工作簿）
except:
    print('优先处理异常：\nExcel中工作簿(sheet)名称必须为：订单信息。请修改Excel后重新运行程序')     #输出提示信息
print("—————————\n抓取工作簿名称：",sheet.title)     #sheet.title为工作簿名称


# 读取XLSX文件  
df = pd.read_excel('myLife.xlsx')  

date_column = '日期'  


# 提取最新日期  
latest_date_str = df[date_column].max()  
# 将字符串转换为日期对象  
print(latest_date_str)
latest_date = datetime.strptime(latest_date_str, '%Y-%m-%d') 
# 提取前面第七天的日期  
seven_days_ago = latest_date - timedelta(days=7)  
  
# 格式化日期为年月日  
latest_date_formatted = latest_date.strftime('%Y-%m-%d')  
seven_days_ago_formatted = seven_days_ago.strftime('%Y-%m-%d')  
  
# 打印结果  
print("最新日期:", latest_date_formatted)  
print("七天前的日期:", seven_days_ago_formatted)

#（1）剔除首行：删除第一行不需要的标题
sheet.delete_rows(1)          #删除行，（）里面数据对应第几行，这里为第一行
print('执行剔除首行成功')     #用来在程序中表现执行成功
# book.save(SavePath)          #save用来保存excel，相当于我们在excel表格中修改后进行另存为。这一步可以最后去做，这里被注释掉了不执行，取消注释进行测试是否删除成功。

#输出工作簿相关参数
print('订单数量：',sheet.max_row)    #程序输出工作簿总共有几行
print('最大列：',sheet.max_column)   #程序输出工作簿总共有几列


Word=Document()    #创建空Word

#全局设置字体
Word.styles['Normal'].font.name=u'宋体'
Word.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#设置页面布局为A4纸张
section=Word.sections[0]
section.page_width = Cm(21) # 设置A4纸的宽度
section.page_height = Cm(29.7) # 设置A4纸的高度
print('导出Word页面的宽度和高度（A4）：', section.page_width.cm,section.page_height.cm)


#首段
str1=Word.add_paragraph(style=None) #增加一个段落
str1_run=str1.add_run('人生导航系统每周报告') #增加文字块
str1_run.bold=True          #加粗
str1_run.font.size=Pt(18)   #行距
str1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#居中

str2=Word.add_paragraph(style=None)
str2_run=str2.add_run(f'{seven_days_ago_formatted}日-{latest_date_formatted}日\n')
str2_run.bold=True
str2_run.font.size=Pt(14)
str2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#居中
  
# 假设 '日期' 列是字符串类型，将其转换为日期时间对象  
df['日期'] = pd.to_datetime(df['日期'])  
  
# 提取前七天的数据  
start_date = seven_days_ago_formatted  
end_date = latest_date_formatted  
df_7days = df[(df['日期'] >= start_date) & (df['日期'] <= end_date)]

# 使用groupby函数统计相同操作的出现次数和总时长  
grouped_df = df_7days.groupby('操作')[['时长']].sum().reset_index() 
grouped_df['出现次数'] = df_7days.groupby('操作').size().reset_index(drop=True) 
grouped_df['数量'] = df_7days.groupby('操作')['数量'].sum().reset_index()['数量'] 
grouped_df['平均数量'] = df_7days.groupby('操作')['数量'].mean().reset_index()['数量'] 
grouped_df['平均时长'] = df_7days.groupby('操作')['时长'].mean().reset_index()['时长'] 
grouped_df['最大时长'] = df_7days.groupby('操作')['时长'].max().reset_index()['时长'] 
grouped_df['最小时长'] = df_7days.groupby('操作')['时长'].min().reset_index()['时长'] 

# # 统计相同操作出现的次数
# operation_counts = df_7days['操作'].value_counts()  
  
# # 统计相同操作对应的总时长  
# total_durations = df_7days.groupby('操作')['时长'].sum()  
  
# # 统计相同操作对应的平均时长  
# average_durations = df_7days.groupby('操作')['时长'].mean()  
  
# # 统计相同操作对应的最大时长  
# max_durations = df_7days.groupby('操作')['时长'].max()  
  
# # 统计相同操作对应的最小时长  
# min_durations = df_7days.groupby('操作')['时长'].min()  
  
# # 将统计结果合并到一个DataFrame中  
# result = pd.concat([operations,operation_counts, total_durations, average_durations, max_durations, min_durations], axis=1)  
# result.columns = ['操作','出现次数', '总时长', '平均时长', '最大时长', '最大时长']  

# # 输出结果  
print(grouped_df)



# 将统计结果添加到Word文档的最上面  
str4=Word.add_paragraph()
str4.paragraph_format.line_spacing=Pt(28)  #行距
str4.paragraph_format.first_line_indent=Cm(1.10)#首行缩进
#一、订单种类分布
str4_run=str4.add_run('以下是本周的操作统计：')
str4_run.font.name=(u'黑体')#字体样式
str4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')#字体样式
str4_run.font.size=Pt(14)#字体大小
str4_run.font.bold=True

str5=Word.add_paragraph()
str5.paragraph_format.line_spacing=Pt(28)  #行距

for index, row in grouped_df.iterrows(): 
    indented_string = f'    本周进行{row["操作"]}活动{row["出现次数"]}次，共计{row["时长"]}分钟,平均{row["平均时长"]}分钟,最长{row["最大时长"]}分钟,最短{row["最小时长"]}分钟 \n'  
      
    str5_run=str5.add_run(indented_string)
    str5_run.font.name=(u'仿宋')#字体样式
    str5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')#字体样式
    str5_run.font.size=Pt(14)#字体大小
    str5_run.bold=True#加粗首段州市文字
    

# 保存文档为.docx文件  
Word.save(f'周报_{seven_days_ago_formatted}-{latest_date_formatted}.docx')

