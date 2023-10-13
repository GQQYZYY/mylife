import pandas as pd
from docx import Document
from datetime import datetime, timedelta

# 读取Excel文件
df = pd.read_excel('myLife.xlsx')

# 确保日期列为日期格式
df['日期'] = pd.to_datetime(df['日期'])

# 获取当前日期并筛选前7天的活动数据
current_date = datetime.now()
start_date = current_date - timedelta(days=7)
end_date = current_date
df_7days = df[(df['日期'] >= start_date) & (df['日期'] <= end_date)]

# 使用pivot_table函数按操作和日期进行透视，并计算总时长和出现次数
pivot_df = pd.pivot_table(df_7days, values='时长', index='操作', columns='日期',fill_value=0, aggfunc='sum')
pivot_df.columns = pivot_df.columns.strftime('%Y-%m-%d')  # 将日期格式化为字符串

# 计算每个操作在每个日期中的时长所占的百分比
pivot_df_percentage = pivot_df.apply(lambda x: x / x.sum(), axis=0)

# 创建新的Word文档
doc = Document()

# 标题
doc.add_heading('以下是本周的操作统计', level=1)

# 创建表格
table = doc.add_table(rows=pivot_df_percentage.shape[0]+2, cols=pivot_df_percentage.shape[1]+1)

# 添加表头
for i, column in enumerate(pivot_df_percentage.columns):
    table.cell(0, i+1).text = column

# 添加行名和数据
for i, index in enumerate(pivot_df_percentage.index):
    table.cell(i+1, 0).text = index
    for j, column in enumerate(pivot_df_percentage.columns):
        table.cell(i+1, j+1).text = '{:.2%}'.format(pivot_df_percentage.loc[index, column])

# 添加总时长行
table.cell(pivot_df_percentage.shape[0]+1, 0).text = '总时长'
for j, column in enumerate(pivot_df_percentage.columns):
    total_duration = pivot_df[column].sum()
    table.cell(pivot_df_percentage.shape[0]+1, j+1).text = str(total_duration)


#可视化
import matplotlib.pyplot as plt
from matplotlib import rcParams  
import os
from docx.shared import Inches
  
# 在这里设置中文字体  
rcParams['font.sans-serif'] = ['SimHei']  # 例如这里我们设置'SimHei'字体，这是一个常用的中文支持的字体  

def pic(l,s,j):
    # 创建饼图  
    fig1, ax1 = plt.subplots() 
    ax1.pie(s, labels=l, autopct='%1.2f%%',  
        shadow=True, startangle=90)  
    # 确保图形是正圆形  
    ax1.axis('equal')    
    # 将图像保存到文件  
    plt.savefig(f'picture{j+1}.png')  # 保存为PNG文件 
    # plt.show()
    
for j in range(7):
    s=[]
    for i in range(len(pivot_df.values)):
        s.append(pivot_df.values[i][j])
    l=pivot_df.index
    pic(l,s,j)
    # 添加一个段落  
    p = doc.add_paragraph(f'{pivot_df.columns[j]}') 
    # # 插入图片，使用os.path.join确保路径格式正确  D:\visual_studio_code\11\mylife
    p.add_run().add_picture(f'picture{j+1}.png', width=Inches(5))


# 格式化文档名称
start_date_str = start_date.strftime('%Y-%m-%d')
end_date_str = end_date.strftime('%Y-%m-%d')

# 保存Word文档
doc.save(f'{start_date_str}_{end_date_str}_activity.docx')