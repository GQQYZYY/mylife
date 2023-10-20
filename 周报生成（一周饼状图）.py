import pandas as pd
import docx
from docx import Document
from datetime import datetime, timedelta
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib.pyplot as plt
from io import BytesIO
import warnings
from docx.shared import Inches
# 读取Excel文件
df = pd.read_excel('myLife.xlsx')

# 确保日期列为日期格式
df['日期'] = pd.to_datetime(df['日期'])

# 获取当前日期并筛选前7天的活动数据
current_date = datetime.now()
start_date = current_date - timedelta(days=7)
end_date = current_date
df_7days = df[(df['日期'] >= start_date) & (df['日期'] <= end_date)]

# 使用pivot_table函数按操作和日期进行透视，并计算总时长和出现次数
pivot_df = pd.pivot_table(df_7days, values='时长', index='操作', columns='日期', aggfunc='sum')
pivot_df.columns = pivot_df.columns.strftime('%Y-%m-%d')  # 将日期格式化为字符串

# 计算每个操作在每个日期中的时长所占的百分比
pivot_df_percentage = pivot_df.apply(lambda x: x / x.sum(), axis=0)

# 统计7天内每天操作为“喝水”的数量
water_data = df_7days[df_7days['操作'] == '喝水']
dates = list(water_data['日期'])
quantities = list(water_data['数量'])

# 统计7天内每天操作为“睡眠”的时长
sleep = df_7days[df_7days['操作'] == '睡眠']
sleep_dates = list(sleep['日期'])
sleep_time = list(sleep['时长'])
# 统计7天内每天操作为“午睡”的时长
noon_break = df_7days[df_7days['操作'] == '午睡']
noon_break_dates = list(noon_break['日期'])
noon_break_time = list(noon_break['时长'])

# 创建新的Word文档
doc = Document()

# 添加标题
title = doc.add_heading('以下是本周的操作统计', level=1)

# 设置标题居中对齐
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# 创建表格
table = doc.add_table(rows=pivot_df_percentage.shape[0]+2, cols=pivot_df_percentage.shape[1]+1)

# 将日期转换为星期几并添加到表头
for i, column in enumerate(pivot_df_percentage.columns):
    date = datetime.strptime(column, '%Y-%m-%d')
    day_of_week = date.strftime('%A')
    table.cell(0, i+1).text = f'{column} ({day_of_week})'

# 将日期转换为星期几并添加到每行数据中
for i, index in enumerate(pivot_df_percentage.index):
    table.cell(i+1, 0).text = index
    for j, column in enumerate(pivot_df_percentage.columns):
        date = datetime.strptime(column, '%Y-%m-%d')
        day_of_week = date.strftime('%A')
        table.cell(i+1, j+1).text = f'{pivot_df_percentage.loc[index, column]:.2%} ({day_of_week})'


# 添加行名和数据
for i, index in enumerate(pivot_df_percentage.index):
    table.cell(i+1, 0).text = index
    for j, column in enumerate(pivot_df_percentage.columns):
        table.cell(i+1, j+1).text = '{:.2%}'.format(pivot_df_percentage.loc[index, column])

# 添加总时长行
table.cell(pivot_df_percentage.shape[0]+1, 0).text = '总时长'
for j, column in enumerate(pivot_df_percentage.columns):
    total_duration = pivot_df[column].sum()
    table.cell(pivot_df_percentage.shape[0]+1, j+1).text = str(total_duration)

# 计算每个操作在7天总时长所占的百分比
total_duration = pivot_df.sum(axis=1)
percentage = total_duration / total_duration.sum()

plt.rcParams['font.sans-serif'] = 'SimSun' # SimSun是一个包含中文字符的字体
# 创建饼图
fig, ax1 = plt.subplots()  # 创建第一个图形框
ax1.pie(percentage, labels=percentage.index, autopct='%1.1f%%', textprops={'fontsize': 8})
ax1.axis('equal')  # 让饼图为圆形
ax1.legend(bbox_to_anchor=(1, 1))  # 显示图例
# 保存饼图
plt.savefig('pie_chart.png')
plt.savefig('pie_chart.png', bbox_inches='tight')
# # 显示饼图
# plt.show()
# 在Word文档中插入饼图
doc.add_picture('pie_chart.png')

# 喝水操作的柱状图
plt.figure()
plt.bar(dates, quantities)
plt.xlabel('日期')
plt.ylabel('数量')
plt.title('7天内每天喝水的数量')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('water.png')
plt.savefig('water.png', bbox_inches='tight')
# plt.show()
# 在Word文档中添加图表
doc.add_picture('water.png')



#睡眠的柱状图 
plt.subplot(1, 2, 1)  # 1行，2列，当前位置在第一个图 
plt.bar(sleep_dates, sleep_time)  
plt.xlabel('日期')  
plt.ylabel('时长 (小时)')  
plt.title('7天内每天睡眠的时长')  
plt.xticks(rotation=45)   
#午睡的柱状图   
plt.subplot(1, 2, 2)  # 1行，2列，当前位置在第二个图  
plt.bar(noon_break_dates, noon_break_time)  
plt.xlabel('日期')  
plt.ylabel('时长 (小时)')    
plt.title('7天内每天午睡的时长')  
plt.xticks(rotation=45)   
plt.tight_layout()  
# 将图像保存为'sleep.png'  
plt.savefig('sleep.png')
doc.add_picture('sleep.png')
# 判断每天的睡眠时长，如果某一天的睡眠时长小于480则输出那一天的睡眠时长不足8小时，请适当调整睡眠时间  
sleep_days = sleep[sleep['时长'] < 480]['日期'].tolist() 
formatted_sleep_days = [day.date().strftime('%Y-%m-%d') for day in sleep_days] 
weekdays = [datetime.strptime(day, '%Y-%m-%d').strftime('%A') for day in formatted_sleep_days]
# 将每个日期转换为星期几的字符串  
# 添加文本到docx文件  
if not weekdays:
    doc.add_paragraph(f'        根据睡眠统计图显示，您的睡眠较为健康，请继续保持！')
else:
    doc.add_paragraph(f'        根据睡眠统计图显示，在{weekdays}，你的睡眠时长不足8个小时，请适当调整休息时间。')

for i in range(len(noon_break_time)):  
    if noon_break_time[i] >= 20:  
        # 添加文本到docx  
        doc.add_paragraph(f'''        小提示：
                午睡时间控制在中午一点到两点期间，20分钟左右可能最为合适。如果是超过一小时的长时间睡眠，它会让人进入睡眠的深睡眠阶段，这个阶段如果醒来的话，就会周身不适、头昏眼花。''')  
        break


# plt.show()
# 格式化文档名称
start_date_str = start_date.strftime('%Y-%m-%d')
end_date_str = end_date.strftime('%Y-%m-%d')

# 保存Word文档
doc.save(f'周报{start_date_str}_{end_date_str}.docx')