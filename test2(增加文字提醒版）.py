import pandas as pd
from docx import Document
from datetime import datetime, timedelta
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Inches

# 读取Excel文件
df = pd.read_excel('myLife.xlsx')

# 确保日期列为日期格式
df['日期'] = pd.to_datetime(df['日期'])

# 获取当前日期并筛选前7天的活动数据
current_date = datetime.now()
start_date = current_date - timedelta(days=7)
end_date = current_date
df_7days = df[(df['日期'] >= start_date) & (df['日期'] <= end_date)]

# 使用pivot_table函数按操作和日期进行透视，并计算总时长和出现次数
pivot_df = pd.pivot_table(df_7days, values='时长', index='操作', columns='日期', aggfunc='sum')
pivot_df.columns = pivot_df.columns.strftime('%Y-%m-%d')  # 将日期格式化为字符串

# 计算每个操作在每个日期中的时长所占的百分比
pivot_df_percentage = pivot_df.apply(lambda x: x / x.sum(), axis=0)

# 创建新的Word文档
doc = Document()
# 添加标题
title = '以下是本周的操作统计'
paragraph = doc.add_paragraph(title)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 添加空行
doc.add_paragraph()

# 创建表格
table = doc.add_table(rows=pivot_df_percentage.shape[0]+2, cols=pivot_df_percentage.shape[1]+1)

# 添加表头
for i, column in enumerate(pivot_df_percentage.columns):
    table.cell(0, i+1).text = column

# 添加行名和数据
for i, index in enumerate(pivot_df_percentage.index):
    table.cell(i+1, 0).text = index
    for j, column in enumerate(pivot_df_percentage.columns):
        table.cell(i+1, j+1).text = '{:.2%}'.format(pivot_df_percentage.loc[index, column])

# 添加总时长行
table.cell(pivot_df_percentage.shape[0]+1, 0).text = '总时长'
for j, column in enumerate(pivot_df_percentage.columns):
    total_duration = pivot_df[column].sum()
    table.cell(pivot_df_percentage.shape[0]+1, j+1).text = str(total_duration)


# 统计本周占时长前五的操作
top5_operations = df_7days.groupby('操作')['时长'].sum().nlargest(5)
top5_operations_list = list(top5_operations.index)

# 添加段落
paragraph = doc.add_paragraph()
paragraph.add_run('本周占时长前五的操作：').bold = True

# 添加列表
for operation in top5_operations_list:
    duration = top5_operations[operation]
    paragraph.add_run('\n\t- {} ({:.2f}小时)'.format(operation, duration)) # 在行前添加 \t 进行缩进




# 按时长对操作进行排序
sorted_operations = df_7days.groupby('操作')['时长'].sum().sort_values(ascending=False)

# 获取前三个操作
top_3_operations = sorted_operations.head(3)

# 检查玩手机是否在前三个操作中
if '玩手机' in top_3_operations.index:
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Inches(0.5)  # 设置首行缩进为0.5英寸
    run = paragraph.add_run('这周可能你在手机上的花费过多的时间，影响了其他重要事务的完成。因此，减少在手机上花费的时间是一个值得考虑的选项。')
    run.bold = True
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('以下是一些具体的建议：')
    run.bold = True
    run = paragraph.add_run('设定时间限制：针对“玩手机”或其他任何不必要的事情，设定一个合理的时间限制。例如，每天在手机上花费的时间不超过1小时或2小时，并逐渐减少时间限制。')
    run = paragraph.add_run('找到替代活动：找一些可以替代玩手机的活动，例如读书、锻炼身体、学习新技能或参与社交活动等。将这些活动纳入你的日常生活中，以减少对手机的依赖。')
    run = paragraph.add_run('坚持时间管理：利用时间管理工具，例如番茄工作法或时间块法等，将工作时间和休息时间合理分配，以提高效率并减少时间的浪费。')
    run = paragraph.add_run('制定奖励机制：为自己设定一些小目标，并在完成后给予自己一些奖励。这种奖励机制可以帮助你保持动力，并减少对手机等不必要事物的依赖。')
    run = paragraph.add_run('通过这些调整日常生活习惯的方法，你可以更好地管理时间并提高效率。将更多的时间用于有意义的活动上，以实现你的目标和梦想。')
    paragraph_format = paragraph.paragraph_format
# 获取所述操作的时间总和
if sum(sorted_operations.loc[['写作业', '编程', '写作', '讨论', '实验', '自学', '预习', '阅读']]) < 2520:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('本周学习时间过短，请适当调整，多花一点时间在学习上。')
    run.bold = True


# 格式化文档名称
start_date_str = start_date.strftime('%Y-%m-%d')
end_date_str = end_date.strftime('%Y-%m-%d')

# 保存Word文档
doc.save(f'{start_date_str}_{end_date_str}周报.docx')