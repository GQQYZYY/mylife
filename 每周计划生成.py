from docx import Document  
from docx.shared import Inches  
from docx.shared import Cm,Inches,Pt   #Word操作：导入单位换算函数
from docx.oxml.ns import qn            #Word操作：中文字体模块
from docx.enum.text import WD_ALIGN_PARAGRAPH     #导入对齐选项
from datetime import datetime, timedelta
import pandas as pd
import random
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
 #读取Excel文件
df = pd.read_excel('myLife.xlsx')

# 确保日期列为日期格式
df['日期'] = pd.to_datetime(df['日期'])
# 创建周报文档  

document = Document() 
#全局设置字体
document.styles['Normal'].font.name=u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

end_date = datetime.now()
start_date = end_date - timedelta(days=7)
last_week_start = start_date - timedelta(days=7)
date_str = start_date.strftime('%Y 年 %m 月 %d 日') + '至 ' + end_date.strftime('%Y 年 %m 月 %d 日')
# 获取上上周的数据
df_last_last_week = df[(df['日期'] >= last_week_start) & (df['日期'] < start_date)]

# 获取上周的数据
df_last_week = df[(df['日期'] >= start_date) & (df['日期'] <= end_date)]
# 找出上周中新增的操作
new_operations = set(df_last_week['操作']) - set(df_last_last_week['操作'])
lost_operations = set(df_last_last_week['操作']) - set(df_last_week['操作'])
# 添加标题  
title = document.add_heading('人生导航周报', level=0)  
title.alignment = 2  
title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#居中

# 添加日期  
title1 = document.add_heading(date_str, level=1)
title1.paragraph_format.line_spacing = Pt(25)


#首段
title2=document.add_paragraph(style=None) #增加一个段落
title2_run=title2.add_run('一.本周内容') #增加文字块
title2_run.bold=True          #加粗
title2_run.font.size=Pt(16)   #字体大小

str2=document.add_paragraph(style=None) #增加一个段落
str2_run=str2.add_run('1、个人成长与技能发展') #增加文字块
str2_run.font.size=Pt(12)
str2_run.bold=True          #加粗
str2.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

str3=document.add_paragraph(style=None) #增加一个段落
str3_run=str3.add_run('（1）完成学习/技能发展目标：') #增加文字块
str3_run.font.size=Pt(12)
str3.paragraph_format.first_line_indent=Cm(2.10)#首行缩进
# 计算阅读操作的总时长
total_reading_time = df_last_week[df_last_week['操作'] == '阅读']['时长'].sum()
str3 = document.add_paragraph(style=None) # 第一部分的内容
str3_run = str3.add_run(f'完成了阅读学习：{total_reading_time}分钟')
str3_run.font.size = Pt(12)
str3.paragraph_format.first_line_indent=Cm(4.10)#首行缩进
# 计算编程操作的总时长
total_programming_time = df_last_week[df_last_week['操作'] == '编程']['时长'].sum()
str4 = document.add_paragraph(style=None) # 新的一行
str4_run = str4.add_run(f'完成了编程学习：{total_programming_time}分钟')
str4_run.font.size = Pt(12)
str4.paragraph_format.first_line_indent = Cm(4.10) # 首行缩进

# 计算预习操作的总时长
total_prep_time = df_last_week[df_last_week['操作'] == '预习']['时长'].sum()
str5 = document.add_paragraph(style=None) # 新的一行
str5_run = str5.add_run(f'进行了知识预习：{total_prep_time}分钟')
str5_run.font.size = Pt(12)
str5.paragraph_format.first_line_indent = Cm(4.10) # 首行缩进

if new_operations:
    str6 = document.add_paragraph(style=None) # 新的一行
    str6_run = str6.add_run('（2）对比上周，本周我做了这些新的事情') # 新段落
    str6_run.font.size = Pt(12)
    str6.paragraph_format.first_line_indent = Cm(2.10) # 首行缩进

    for operation in new_operations:
        new_operation_paragraph = document.add_paragraph(style=None)
        new_operation_run = new_operation_paragraph.add_run(f'{operation} 时长：{df_last_week[df_last_week["操作"] == operation]["时长"].sum()}分钟')
        new_operation_run.font.size = Pt(12)
        new_operation_paragraph.paragraph_format.first_line_indent = Cm(4.10) # 首行缩进
#健康与健身
str7=document.add_paragraph(style=None) #增加一个段落
str7_run=str7.add_run('2.健康与健身') #增加文字块
str7_run.font.size=Pt(12)
str7_run.bold=True          #加粗
str7.paragraph_format.first_line_indent=Cm(1.10)#首行缩进
#饮食
str8=document.add_paragraph(style=None) #增加一个段落
str8_run=str8.add_run('（1）饮食：') #增加文字块
str8_run.font.size=Pt(12)
str8.paragraph_format.first_line_indent=Cm(2.10)#首行缩进

# 计算上一周吃午餐和吃晚餐时长超过 30 分钟的情况
lunch_duration = df_last_week[(df_last_week['操作'].str.contains('吃午餐')) & (df_last_week['时长'] > 30)]['时长'].sum()
dinner_duration = df_last_week[(df_last_week['操作'].str.contains('吃晚餐')) & (df_last_week['时长'] > 30)]['时长'].sum()
# 统计上周“吃早餐”的次数
breakfast_count = (df_last_week['操作'] == '吃早餐').sum()

if lunch_duration > 0 or dinner_duration > 0:
    str9 = document.add_paragraph(style=None)
    str9_run = str9.add_run('请注意午餐或晚餐时间不宜过长，因为过长的进食时间可能导致过度摄入食物或使消化不良，从而对健康产生负面影响。')
    str9_run.font.size = Pt(12)
    str9.paragraph_format.first_line_indent = Cm(4.10) # 首行缩进
# 根据计数提供消息
if breakfast_count < 7:
    str10 = document.add_paragraph(style=None)
    str10_run = str10.add_run(f'上周你只吃了 {breakfast_count} 次早餐，请确保每天都吃早餐，因为早餐是一天中最重要的一餐。')
    str10_run.font.size = Pt(12)
    str10.paragraph_format.first_line_indent = Cm(4.10)
if breakfast_count == 7 and lunch_duration == 0 and dinner_duration == 0:
    str11 = document.add_paragraph(style=None)
    str11_run = str11.add_run('本周你已坚持健康的饮食习惯，请继续坚持并且注重营养均衡哦。')
    str11_run.font.size = Pt(12)
#运动
str8=document.add_paragraph(style=None) #增加一个段落
str8_run=str8.add_run('（2）运动：') #增加文字块
str8_run.font.size=Pt(12)
str8.paragraph_format.first_line_indent=Cm(2.10)#首行缩进
# 初始化变量来保存步行的总时长和步行的数量总值
total_walking_time = df_last_week[df_last_week['操作'] == '步行']['时长'].sum()
total_walking_count = df_last_week[df_last_week['操作'] == '步行']['数量'].sum()

# 输出步行的总时长和步行的数量总值
str9 = document.add_paragraph(style=None)
str9_run = str9.add_run(f'上周步行总时长：{total_walking_time}分钟，共步行{total_walking_count}步。')
str9_run.font.size = Pt(12)
str9.paragraph_format.first_line_indent = Cm(4.10)

#财务与理财分类
str7=document.add_paragraph(style=None) #增加一个段落
str7_run=str7.add_run('3.财务与理财') #增加文字块
str7_run.font.size=Pt(12)
str7_run.bold=True          #加粗
str7.paragraph_format.first_line_indent=Cm(1.10)#首行缩进
#充值活动
str8=document.add_paragraph(style=None) #增加一个段落
str8_run=str8.add_run('（1）充值/购物：') #增加文字块
str8_run.font.size=Pt(12)
str8.paragraph_format.first_line_indent=Cm(2.10)#首行缩进
# 检查是否存在充值操作的记录
if '充值' in df_last_week['操作'].values:
    total_recharge_count = df_last_week[df_last_week['操作'].str.contains('充值')]['数量'].sum()
else:
    total_recharge_count = 0

# 统计上一周操作购物的总时长和总数量
total_shopping_duration = df_last_week[df_last_week['操作'] == '购物']['时长'].sum()
total_shopping_count = df_last_week[df_last_week['操作'] == '购物']['数量'].sum()

if total_recharge_count > 0:
    str9 = document.add_paragraph(style=None)
    str9_run = str9.add_run(f'上周充值总金额为：{total_recharge_count}元。')
    str9_run.font.size = Pt(12)
    str9.paragraph_format.first_line_indent = Cm(4.10)

str10 = document.add_paragraph(style=None)
str10_run = str10.add_run(f'上周购物总时长：{total_shopping_duration}分钟，共消费{total_shopping_count}元。')
str10_run.font.size = Pt(12)
str10.paragraph_format.first_line_indent = Cm(4.10)

#4.社交分类
str7=document.add_paragraph(style=None) #增加一个段落
str7_run=str7.add_run('4.社交') #增加文字块
str7_run.font.size=Pt(12)
str7_run.bold=True          #加粗
str7.paragraph_format.first_line_indent=Cm(1.10)#首行缩进
# 统计上一周操作网聊、讨论和购物的总时长和总数量
total_online_chat_duration = df_last_week[df_last_week['操作'] == '网聊']['时长'].sum()
total_online_chat_count = df_last_week[df_last_week['操作'] == '网聊']['数量'].sum()

total_discussion_duration = df_last_week[df_last_week['操作'] == '讨论']['时长'].sum()
total_discussion_count = df_last_week[df_last_week['操作'] == '讨论']['数量'].sum()

if total_online_chat_duration > 0 or total_discussion_duration > 0 or total_shopping_duration > 0:
    str8 = document.add_paragraph(style=None)
    str8_run = str8.add_run('（1）与朋友的关系：')
    str8_run.font.size = Pt(12)
    str8.paragraph_format.first_line_indent = Cm(2.10)

    if total_online_chat_duration > 0:
        str8 = document.add_paragraph(style=None)
        str8_run = str8.add_run(f'上周你和你的朋友们进行了总时长：{total_online_chat_duration}分钟的网聊。')
        str8_run.font.size = Pt(12)
        str8.paragraph_format.first_line_indent = Cm(4.10)

    if total_discussion_duration > 0:
        str9 = document.add_paragraph(style=None)
        str9_run = str9.add_run(f'上周你和你的朋友们进行了总时长：{total_discussion_duration}分钟的讨论。')
        str9_run.font.size = Pt(12)
        str9.paragraph_format.first_line_indent = Cm(4.10)

    if total_shopping_duration > 0:
        str10 = document.add_paragraph(style=None)
        str10_run = str10.add_run(f'上周你和你的朋友们进行了总时长：{total_shopping_duration}分钟的购物活动，共购物{total_shopping_count}元。')
        str10_run.font.size = Pt(12)
        str10.paragraph_format.first_line_indent = Cm(4.10)

#5.休闲与娱乐
str7=document.add_paragraph(style=None) #增加一个段落
str7_run=str7.add_run('5.休闲与娱乐') #增加文字块
str7_run.font.size=Pt(12)
str7_run.bold=True          #加粗
str7.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

str8=document.add_paragraph(style=None) #增加一个段落
str8_run=str8.add_run('（1）活动与兴趣爱好：') #增加文字块
str8_run.font.size=Pt(12)
str8.paragraph_format.first_line_indent=Cm(2.10)#首行缩进

excluded_activities = ["吃早餐", "吃午餐", "吃晚餐", "排泄", "喝水", "洗漱", "睡眠", "午睡", "思考", "步行", "值班", "预习", "自学", "讨论", "玩手机","写作业","网聊","休养","充值","分发","上课","洗碗"]

# 清理字符串并找出上一周操作中出现5次或更多的活动，不在排除列表中
frequent_activities = df_last_week['操作'].str.strip().value_counts()
frequent_activities = frequent_activities[(frequent_activities >= 5) & (~frequent_activities.index.isin(excluded_activities))]

# 创建一个新的段落来展示这些活动
str11 = document.add_paragraph(style=None)
str11_run = str11.add_run('根据对上周的统计分析以下活动也许是你的爱好或者你经常做的活动，如果是你喜欢的请继续坚持下去：')
str11_run.font.size = Pt(12)
str11.paragraph_format.first_line_indent = Cm(4.10)

# 遍历这些活动并展示它们的时长
for activity, count in frequent_activities.items():
    total_duration = df_last_week[df_last_week['操作'] == activity]['时长'].sum()
    activity_paragraph = document.add_paragraph(style=None)
    activity_run = activity_paragraph.add_run(f'{activity} 总时长：{total_duration}分钟。')
    activity_run.font.size = Pt(12)
    activity_paragraph.paragraph_format.first_line_indent = Cm(4.10)
#二、总结与反馈
title3=document.add_paragraph(style=None) #增加一个段落
title3_run=title3.add_run('二.总结与反馈') #增加文字块
title3_run.bold=True          #加粗
title3.paragraph_format.line_spacing=Pt(30)  #行距
title3_run.font.size=Pt(20)   #字体大小


study=['编程','预习','阅读']
health=['步行']
social=['网聊','讨论','购物']
spend=['充值','购物']

# 求各类操作的总时长
def time_activity(set1,df_last_week):
    # set2为df_last_week
    x=0
    for i in range(len(set1)):
        for j in range(len(df_last_week['操作'])):
            if set1[i] == df_last_week['操作'].values[j]:
                x=x+df_last_week['时长'].values[j]     
    return x

# 求某个操作的对象集
def object_activity(value,df_last_week):
    y=[]
    for i in range(len(df_last_week)):
        if value == df_last_week['操作'].values[i]:
            y.append(df_last_week['对象'].values[i])
    
    y=list(set(y))
    y='、'.join(y)
    return y

#首段
# str2=document.add_paragraph(style=None) #增加一个段落
# str2_run=str2.add_run('1、活动总结') #增加文字块
# str2_run.font.size=Pt(12)
# str2_run.bold=True          #加粗
# str2.paragraph_format.first_line_indent=Cm(1.10)#首行缩进



# 学习总结
xx=time_activity(study,df_last_last_week)
str3=document.add_paragraph(style=None) #增加一个段落
str3_run=str3.add_run(f'（1）上课之余，上周花费了{xx}分钟在学习上，约每天学习{int(xx/420)}了小时。建议每天学习8小时，不超过10小时，也要劳逸结合。') #增加文字块
str3_run.font.size=Pt(12)
str3.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

# 健康生活总结
js=time_activity(health,df_last_last_week)
str3=document.add_paragraph(style=None) #增加一个段落
str3_run=str3.add_run(f'（2）上周健身共计{js}分钟，相当于每天运动{int(js/7)}了分钟，请继续保持运动，健康生活。') #增加文字块
str3_run.font.size=Pt(12)
str3.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

# 社交总结
sj=time_activity(social,df_last_last_week)
str3=document.add_paragraph(style=None) #增加一个段落
str3_run=str3.add_run(f'（3）在社交活动上用了{sj}分钟，社交可以保持人际关系、提升自我认知，努力成为交际小达人吧。') #增加文字块
str3_run.font.size=Pt(12)
str3.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

# 消费总结
# total_shopping_count=50
sj=time_activity(social,df_last_last_week)
str3=document.add_paragraph(style=None) #增加一个段落
if total_shopping_count >= 1000:
    str3_run=str3.add_run(f'（4）上周花费了{total_shopping_count}元，要注意保证日常消费支出，不要大手大脚。') #增加文字块
elif total_shopping_count <= 100:
    str3_run=str3.add_run(f'（4）上周花费了{total_shopping_count}元,恭喜成为省钱小能手。') #增加文字块
else:
    str3_run=str3.add_run(f'（4）上周花费了{total_shopping_count}元,支出水平正常。') #增加文字块
str3_run.font.size=Pt(12)
str3.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

# 兴趣总结
xq=list(frequent_activities.index)
sj=time_activity(social,df_last_last_week)
str3=document.add_paragraph(style=None) #增加一个段落
if int(time_activity(xq,df_last_week)/420) >= 6:
    str3_run=str3.add_run(f'（5）上周扩展了{len(xq)}项兴趣爱好，平均每天投入时间{int(time_activity(xq,df_last_week)/7)}分钟。培养兴趣爱好有很多好处，但如果沉迷于兴趣爱好，也会对生活和健康产生负面影响，注意平衡兴趣爱好和生活') #增加文字块
else:
    str3_run=str3.add_run(f'（5）上周扩展了{len(xq)}项兴趣爱好，平均每天投入时间{int(time_activity(xq,df_last_week)/7)}分钟。培养兴趣爱好可以使心情愉悦，提升自我能力，加油。') #增加文字块
str3_run.font.size=Pt(12)
str3.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

# # 上课总结
# kc=object_activity('上课',df_last_week)
# str3=document.add_paragraph(style=None) #增加一个段落
# str3_run=str3.add_run(f'（6）上周学习了{kc}课程，记得复习巩固。') #增加文字块
# str3_run.font.size=Pt(12)
# str3.paragraph_format.first_line_indent=Cm(1.10)#首行缩进

# # 阅读总结
# yd=object_activity('阅读',df_last_week)
# if len(yd) != 0:
#     str3=document.add_paragraph(style=None) #增加一个段落
#     str3_run=str3.add_run(f'（7）读书使人进步。上周阅读了{yd}等书籍。通过读书，我们可以了解到不同的人物和故事，从而更好地理解和处理自己的情感。') #增加文字块
#     str3_run.font.size=Pt(12)
#     str3.paragraph_format.first_line_indent=Cm(2.10)#首行缩进


# #三、下周计划
# 获取历史操作的最后30天内容
last_30_days = df["操作"].tail(500)
# 去重操作
unique_operations = last_30_days.unique()
# 要排除的操作名称
exclude_operations = ["吃早餐", "吃午餐", "吃晚餐", "排泄", "喝水", "洗漱", "睡眠", "午睡", "思考", "记录"]
# 从历史操作中随机选择5到10个不同的操作，排除指定的操作
available_operations = [operation for operation in unique_operations if operation not in exclude_operations]
if len(available_operations) >= 5:
    num_operations = random.randint(5, min(10, len(available_operations)))
    random_operations = random.sample(available_operations, num_operations)

title4=document.add_paragraph(style=None) #增加一个段落
title4_run=title4.add_run('三.下周计划') #增加文字块
title4_run.bold=True          #加粗
title4.paragraph_format.line_spacing=Pt(30)  #行距
title4_run.font.size=Pt(20)   #字体大小

str4_1 = document.add_paragraph(style=None)
str4_1_run = str4_1.add_run(f'1.活动预测')
str4_1_run.font.size = Pt(12)
str4_1.paragraph_format.first_line_indent = Cm(1.10)
str4_1_run.bold=True          #加粗

str4_2 = document.add_paragraph(style=None)
str4_2_run = str4_2.add_run(f'（1）根据前两周的活动内容，预测你下周可能会进行的活动：')
str4_2_run.font.size = Pt(12)
str4_2.paragraph_format.first_line_indent = Cm(2.10)
# # 将随机生成的操作内容用逗号分隔开，不换行
operations_text = "、".join(random_operations)
p1 = document.add_paragraph(operations_text)
for run in p1.runs:
    run.font.size = Pt(12)
    run.bold=True
    p1.paragraph_format.first_line_indent = Cm(0)
# # 添加上周新增的操作到文档
if new_operations:
    new_operations_text = "（2）对比上周，本周你尝试了"
    p = document.add_paragraph()
    # 添加非加粗文本
    p.add_run(new_operations_text)
#    # 添加加粗文本，并使用逗号分隔
    for i, operation in enumerate(new_operations):
        if i > 0:
            # 在文本之间添加逗号
            p.add_run("、")
        # 添加加粗文本
        run = p.add_run(operation)
        run.bold = True
    new_operations_text = "活动，在接下来的一周你也可能会进行这些活动。如果你觉得这个操作是有利于你的，请继续保持。"
#     # 添加非加粗文本
    p.add_run(new_operations_text)
    p.paragraph_format.first_line_indent = Cm(2.01)  # 设置首行缩进

# # 添加上周失去的操作到文档
if lost_operations:
    lost_operations_text = "（3）与上周相比，本周你停止了"
    p_lost = document.add_paragraph()
    
#     # 添加非加粗文本
    p_lost.add_run(lost_operations_text)

#     # 添加加粗文本，并使用逗号分隔
    for i, operation in enumerate(lost_operations):
        if i > 0:
            # 在文本之间添加逗号
            p_lost.add_run("、")
#         # 添加加粗文本
        run = p_lost.add_run(operation)
        run.bold = True

    lost_operations_text = "活动，如果这是因为你觉得这是不好的事情，那你太棒了！但如果是由于你每周计划没有按计划完成，那么请合理安排时间，下周也请继续完成计划。"
    
#     # 添加非加粗文本
    p_lost.add_run(lost_operations_text)

    p_lost.paragraph_format.first_line_indent = Cm(2.01)  # 设置首行缩进

# 格式化文档名称
start_date_str = start_date.strftime('%Y-%m-%d')
end_date_str = end_date.strftime('%Y-%m-%d')

# 保存Word文档
document.save(f'{start_date_str}_{end_date_str}周报.docx')